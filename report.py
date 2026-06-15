"""
report.py
---------
Renders the Facility Assessment Snapshot as a polished PDF (fpdf2) and an
editable Word document (python-docx).

Single source of truth: `build_report_rows()` produces the ordered (label, value)
list used by BOTH renderers, so the PDF and the .docx can never drift apart.

Branding guardrail (explicit requirement): the banner string
"INFINITE — Managed by MEDELITE" is hard-coded and is NEVER replaced by the
facility name. The facility name lives only in the body row "Name of Facility".
"""

from __future__ import annotations

import io
import os

from fpdf import FPDF

from cms_client import fmt_metric

# --- brand constants (hard-coded; never overwritten by API/user data) --------
BRAND_PLATFORM = "INFINITE"
BRAND_LINE = "INFINITE — Managed by MEDELITE"   # em dash, exact per brief
REPORT_TITLE = "FACILITY ASSESSMENT SNAPSHOT"
MAGENTA = (214, 0, 126)       # INFINITE brand
BLUE = (57, 147, 203)         # MedElite #3993CB
DEEP = (6, 106, 171)          # MedElite #066AAB
DARK = (33, 37, 41)
TABLE_HEADER = (57, 147, 203)
SECTION_BG = (230, 241, 249)
SECTION_TX = (6, 106, 171)
ZEBRA = (246, 250, 253)
VALUE_TX = (70, 80, 90)
BORDER = (188, 210, 228)   # visible grid / column divider

_FONT_DIR = os.path.join(os.path.dirname(__file__), "assets", "fonts")

# Maps the 12 hospitalization labels -> (claims key, value slot, fmt kind)
_HOSP_LAYOUT = [
    ("Short Term Hospitalization",                  ("STR", "hosp"), "facility", "STR_PCT"),
    ("STR National Avg. for Hospitalization",       ("STR", "hosp"), "nation",   "STR_PCT"),
    ("STR State National Avg. for Hospitalization", ("STR", "hosp"), "state",    "STR_PCT"),
    ("STR ED Visit",                                ("STR", "ed"),   "facility", "STR_PCT"),
    ("STR ED Visits National Avg.",                 ("STR", "ed"),   "nation",   "STR_PCT"),
    ("STR ED Visits State Avg.",                    ("STR", "ed"),   "state",    "STR_PCT"),
    ("LT Hospitalization",                          ("LT", "hosp"),  "facility", "LT"),
    ("LT National Avg. for Hospitalization",        ("LT", "hosp"),  "nation",   "LT"),
    ("LT State National Avg. for Hospitalization",  ("LT", "hosp"),  "state",    "LT"),
    ("ED Visit",                                    ("LT", "ed"),    "facility", "LT"),
    ("LT ED Visits National Avg.",                  ("LT", "ed"),    "nation",   "LT"),
    ("LT ED Visits State Avg.",                     ("LT", "ed"),    "state",    "LT"),
]


def resolve_display_name(snap: dict, manual: dict) -> str:
    """API legal name by default; user's custom name overrides it if provided."""
    override = (manual.get("name_override") or "").strip()
    return override if override else snap.get("legal_name", "")


def _dash(v) -> str:
    """Show an em dash for blank fields so the report reads as intentional."""
    s = "" if v is None else str(v)
    return s if s.strip() else "—"


def build_report_rows(snap: dict, manual: dict) -> list[tuple[str, str]]:
    """Ordered (label, value) rows — the single source of truth for both exports."""
    rows = [
        ("Name of Facility", resolve_display_name(snap, manual)),
        ("Location", snap.get("address", "")),
        ("EMR", manual.get("emr", "")),
        ("Census Capacity", str(snap.get("certified_beds", "") or "")),
        ("Current Census", str(manual.get("current_census", "") or "")),
        ("Type of Patient", manual.get("patient_type", "")),
        ("Previous Coverage from Medelite", manual.get("prev_coverage", "")),
        ("Previous Provider Performance from Medelite", manual.get("prev_performance", "")),
        ("Medical Coverage", manual.get("medical_coverage", "")),
        ("Overall Star Rating", str(snap["ratings"].get("overall", "") or "")),
        ("Health Inspection", str(snap["ratings"].get("health_inspection", "") or "")),
        ("Staffing", str(snap["ratings"].get("staffing", "") or "")),
        ("Quality of Resident Care", str(snap["ratings"].get("quality", "") or "")),
    ]
    rows = [(label, _dash(value)) for label, value in rows]
    hosp = snap.get("hospitalization")
    if hosp:
        for label, key, slot, kind in _HOSP_LAYOUT:
            rows.append((label, fmt_metric(hosp[key][slot], kind)))
    return rows


# --- PDF ---------------------------------------------------------------------
class _SnapshotPDF(FPDF):
    def __init__(self, state_abbr: str):
        super().__init__(orientation="P", unit="mm", format="letter")
        self.state_abbr = state_abbr
        self._has_unicode = self._register_fonts()
        self.set_auto_page_break(auto=True, margin=18)
        self.set_margins(18, 16, 18)

    def _register_fonts(self) -> bool:
        # Carlito: open-source, metrically identical to Calibri (legal to bundle
        # in a public repo and present on the Linux deploy host, unlike Calibri).
        try:
            self.add_font("Calibri", "", os.path.join(_FONT_DIR, "Carlito-Regular.ttf"))
            self.add_font("Calibri", "B", os.path.join(_FONT_DIR, "Carlito-Bold.ttf"))
            self.font_family_name = "Calibri"
            return True
        except Exception:
            self.font_family_name = "Helvetica"   # graceful fallback
            return False

    def header(self):
        # --- hard-coded brand banner (guardrail) ---
        if self._has_unicode:
            self.set_font(self.font_family_name, "B", 20)
            self.set_text_color(*MAGENTA)
            self.cell(0, 10, "INFINITE", align="C", new_x="LMARGIN", new_y="NEXT")
            self.set_font(self.font_family_name, "B", 11)
            self.set_text_color(*BLUE)
            self.cell(0, 6, "Managed by MEDELITE", align="C",
                      new_x="LMARGIN", new_y="NEXT")
        else:
            self.set_font(self.font_family_name, "B", 16)
            self.set_text_color(*MAGENTA)
            self.cell(0, 9, "INFINITE - Managed by MEDELITE", align="C",
                      new_x="LMARGIN", new_y="NEXT")
        self.ln(2)
        self.set_draw_color(*BLUE)
        self.set_line_width(0.5)
        self.line(18, self.get_y(), self.w - 18, self.get_y())
        self.ln(3)
        self.set_font(self.font_family_name, "B", 13)
        self.set_text_color(*DARK)
        self.cell(0, 7, REPORT_TITLE, align="C", new_x="LMARGIN", new_y="NEXT")
        self.set_font(self.font_family_name, "B", 11)
        self.cell(0, 6, self.state_abbr, align="C", new_x="LMARGIN", new_y="NEXT")
        self.ln(3)


def render_pdf(snap: dict, manual: dict) -> bytes:
    from datetime import datetime

    rows = build_report_rows(snap, manual)
    pdf = _SnapshotPDF(snap.get("state", ""))
    # Proper document metadata: makes the file a legitimate, identifiable PDF
    # (also reduces heuristic false-positives from aggressive AV download scanners).
    disp = resolve_display_name(snap, manual)
    pdf.set_title(f"Facility Assessment Snapshot - {disp}")
    pdf.set_author("INFINITE - Managed by MEDELITE")
    pdf.set_subject("CMS skilled nursing facility assessment report")
    pdf.set_creator("Facility Assessment Snapshot")
    pdf.set_producer("Facility Assessment Snapshot (fpdf2)")
    pdf.set_creation_date(datetime.now())
    pdf.add_page()
    fam = pdf.font_family_name

    x0 = pdf.l_margin
    label_w = 96
    value_w = pdf.w - pdf.l_margin - pdf.r_margin - label_w
    full_w = label_w + value_w
    row_h, sec_h = 6.5, 5.7
    sections = {0: "Facility Profile", 9: "CMS Star Ratings",
                13: "Hospitalization & ED Metrics"}
    pdf.set_draw_color(*BORDER)
    pdf.set_line_width(0.25)

    # column header (boxed; right border on Field = the column divider)
    pdf.set_x(x0)
    pdf.set_font(fam, "B", 9.5)
    pdf.set_fill_color(*TABLE_HEADER)
    pdf.set_text_color(255, 255, 255)
    pdf.cell(label_w, row_h + 0.6, "  Field", border="LTRB", fill=True,
             new_x="RIGHT", new_y="TOP")
    pdf.cell(value_w, row_h + 0.6, "Value", border="TRB", fill=True,
             new_x="LMARGIN", new_y="NEXT")

    for i, (label, value) in enumerate(rows):
        if i in sections:
            pdf.set_x(x0)
            pdf.set_font(fam, "B", 8)
            pdf.set_fill_color(*SECTION_BG)
            pdf.set_text_color(*SECTION_TX)
            pdf.cell(full_w, sec_h, "  " + sections[i].upper(), border="LRB",
                     fill=True, new_x="LMARGIN", new_y="NEXT")
        pdf.set_fill_color(*(ZEBRA if i % 2 == 0 else (255, 255, 255)))
        y0 = pdf.get_y()
        pdf.set_x(x0)
        pdf.set_font(fam, "B", 9)
        pdf.set_text_color(*DARK)
        # label cell: left outer + right divider + bottom
        pdf.multi_cell(label_w, row_h, "  " + label, border="LRB", align="L",
                       fill=True, new_x="RIGHT", new_y="TOP", max_line_height=4)
        pdf.set_xy(x0 + label_w, y0)
        pdf.set_font(fam, "", 9)
        pdf.set_text_color(*VALUE_TX)
        # value cell: right outer + bottom
        pdf.multi_cell(value_w, row_h, str(value), border="RB", align="L",
                       fill=True, new_x="LMARGIN", new_y="NEXT", max_line_height=4)

    # --- clickable Medicare source hyperlink (dynamic CCN) ---
    pdf.ln(4)
    url = snap.get("care_compare_url", "")
    pdf.set_font(fam, "B", 9.5)
    pdf.set_text_color(*DEEP)
    pdf.cell(0, 6, "View official CMS Care Compare profile  >", align="L",
             link=url, new_x="LMARGIN", new_y="NEXT")
    pdf.set_font(fam, "", 7.5)
    pdf.set_text_color(120, 120, 120)
    proc = snap.get("processing_date", "")
    stamp = f"Data as of {proc} | Source: CMS Provider Data Catalog (data.cms.gov)"
    pdf.cell(0, 5, stamp, align="L", new_x="LMARGIN", new_y="NEXT")

    return _finalize_pdf(bytes(pdf.output()))


def _finalize_pdf(pdf_bytes: bytes) -> bytes:
    """Re-emit the PDF without the (benign) document-level /OpenAction.

    fpdf2 always writes an /OpenAction that sets the initial view/zoom. A static
    report needs no auto-action on open, and removing it keeps the file a minimal,
    purely-static document — which avoids the heuristic flag some AV download
    scanners raise on freshly generated files. The clickable Care Compare
    hyperlink and document metadata are preserved. Falls back to the original
    bytes if anything goes wrong, so a download never fails.
    """
    try:
        from pypdf import PdfReader, PdfWriter
        reader = PdfReader(io.BytesIO(pdf_bytes))
        writer = PdfWriter()
        for page in reader.pages:
            writer.add_page(page)
        if reader.metadata:
            writer.add_metadata({str(k): str(v) for k, v in reader.metadata.items()})
        buf = io.BytesIO()
        writer.write(buf)
        return buf.getvalue()
    except Exception:
        return pdf_bytes


# --- Word (.docx) ------------------------------------------------------------
def render_docx(snap: dict, manual: dict) -> bytes:
    import docx
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    rows = build_report_rows(snap, manual)
    d = docx.Document()
    d.styles["Normal"].font.name = "Calibri"  # Word ships Calibri natively

    def _center(p):
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        return p

    # --- brand banner (guardrail: hard-coded) ---
    p = _center(d.add_paragraph())
    r = p.add_run("INFINITE")
    r.bold = True; r.font.size = Pt(22); r.font.color.rgb = RGBColor(*MAGENTA)
    p2 = _center(d.add_paragraph())
    r2 = p2.add_run("Managed by MEDELITE")
    r2.bold = True; r2.font.size = Pt(12); r2.font.color.rgb = RGBColor(*BLUE)
    t = _center(d.add_paragraph()); rt = t.add_run(REPORT_TITLE)
    rt.bold = True; rt.font.size = Pt(14)
    st = _center(d.add_paragraph()); rs = st.add_run(snap.get("state", ""))
    rs.bold = True; rs.font.size = Pt(12)

    table = d.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    for label, value in rows:
        c = table.add_row().cells
        rl = c[0].paragraphs[0].add_run(label); rl.bold = True; rl.font.size = Pt(9.5)
        rv = c[1].paragraphs[0].add_run(str(value)); rv.font.size = Pt(9.5)

    # clickable hyperlink paragraph
    d.add_paragraph()
    link_p = d.add_paragraph()
    _add_hyperlink(link_p, snap.get("care_compare_url", ""),
                   "View official CMS Care Compare profile →", BLUE)
    foot = d.add_paragraph()
    fr = foot.add_run(
        f"Data as of {snap.get('processing_date','')} · "
        "Source: CMS Provider Data Catalog (data.cms.gov)")
    fr.font.size = Pt(7.5); fr.font.color.rgb = RGBColor(120, 120, 120)

    buf = io.BytesIO(); d.save(buf)
    return buf.getvalue()


def _add_hyperlink(paragraph, url: str, text: str, rgb):
    """Insert a real clickable hyperlink into a python-docx paragraph."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "%02X%02X%02X" % rgb)
    u = OxmlElement("w:u"); u.set(qn("w:val"), "single")
    b = OxmlElement("w:b")
    rpr.append(b); rpr.append(color); rpr.append(u)
    new_run.append(rpr)
    t = OxmlElement("w:t"); t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink
